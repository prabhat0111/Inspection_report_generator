from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx2pdf import convert
import os

def add_section_heading(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(178, 34, 34)
    para.space_after = Pt(8)
    return para

def add_normal_paragraph(doc, text):
    para = doc.add_paragraph(text)
    para.space_after = Pt(6)
    return para

def generate_report_and_pdf(data, photo_folder, report_folder):
    doc = Document("template.docx")

    add_section_heading(doc, "FIRST INSPECTION REPORT")
    add_normal_paragraph(doc, f"INSURED/POLICYHOLDER: {data['insured_name']}")
    add_normal_paragraph(doc, f"ADDRESS: {data['address']}")
    add_normal_paragraph(doc, f"INSURER: {data['insurer']}")
    add_normal_paragraph(doc, f"CLAIM #: {data['claim_number']}")
    add_normal_paragraph(doc, "ADJUSTER/ CLAIM REP: TOP GUN")
    add_normal_paragraph(doc, f"DATE OF INSPECTION: {data['date_of_inspection']}")
    add_normal_paragraph(doc, f"DATE OF LOSS: {data['date_of_loss']}")
    add_normal_paragraph(doc, f"DATE OF REPORT: {data['date_of_report']}")
    add_normal_paragraph(doc, f"TYPE OF LOSS: {data['type_of_loss']}")

    # Fron
    front_photo = next((f for f in os.listdir(photo_folder) if 'front' in f.lower()), None)
    if front_photo:
        doc.add_picture(os.path.join(photo_folder, front_photo), width=Inches(4.5))

    add_section_heading(doc, "\nCAUSE OF LOSS:")
    add_normal_paragraph(doc, data['cause_of_loss'])

    add_section_heading(doc, "\nSCOPE OF WORK:")
    add_normal_paragraph(doc, "1. Assess, pack and move out all salvageable contents.\n"
                              "2. Inventory all the affected contents.\n"
                              "3. Inspect all affected electronics.\n"
                              "4. Restore salvageable contents.\n"
                              "5. Dispose of non-salvageable contents.")

    add_section_heading(doc, "\nRECOMMENDED RESERVES FOR TRINITY’S INVOLVEMENT:")
    add_normal_paragraph(doc, f"• Indemnity Work : Should not exceed ${data['indemnity_work']} plus HST.")
    add_normal_paragraph(doc, f"• Trinity Listing & Pricing Expense Reserve: Should not exceed ${data['listing_pricing_reserve']} plus HST")

    add_section_heading(doc, "\nRECOMMENDED RESERVES FOR THE TOTAL CONTENTS LOSS:")
    add_normal_paragraph(doc, f"Based on a visual inspection of the extent of non-salvageable items on the main floor, we believe that the total replacement cost for the non-salvageable items should not exceed ${data['contents_loss_reserve']} plus HST.")

    add_section_heading(doc, "\nCONCLUSION:")
    add_normal_paragraph(doc, "Once our scope of work is approved, we can attend and begin the pack out process.")

    add_normal_paragraph(doc, "\nThank You,\n\nMo Waez\nTrinity Contents Management\nmo@trinitycontents.com\n(647) 613-2246")

    
    section_keywords = {
        "KITCHEN & DINING AREA": ["kitchen", "dining"],
        "LIVING ROOM": ["living"],
        "BEDROOM_A": ["bedroom_a"],
        "BEDROOM_B": ["bedroom_b"],
        "STORAGE ROOM": ["storage", "storeroom"]
    }

    section_photos = {section: [] for section in section_keywords}
    for img_file in sorted(os.listdir(photo_folder)):
        lower_name = img_file.lower()
        for section, keywords in section_keywords.items():
            if any(keyword in lower_name for keyword in keywords):
                section_photos[section].append(os.path.join(photo_folder, img_file))
                break

    for section, images in section_photos.items():
        if images:
            doc.add_page_break()
            add_section_heading(doc, section)
            for img_path in images:
                doc.add_picture(img_path, width=Inches(3.5))
                doc.add_paragraph()

    word_path = os.path.join(report_folder, f"{data['claim_number']}.docx")
    doc.save(word_path)

    pdf_path = os.path.join(report_folder, f"{data['claim_number']}.pdf")
    convert(word_path, pdf_path)

    return word_path, pdf_path
